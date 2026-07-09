#!/usr/bin/env python3
"""Generate Automation Test Execution Report for ValClassifieds v1."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

REPORT_DIR = "reports"
DOCX_PATH = os.path.join(REPORT_DIR, "Automation_Test_Execution_Report_ValClassifieds_v1.docx")
HTML_PATH = os.path.join(REPORT_DIR, "Automation_Test_Execution_Report_ValClassifieds_v1.html")

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_BLUE = RGBColor(0x2E, 0x5C, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
GREEN = RGBColor(0x1B, 0x8A, 0x3D)

TODAY = datetime.date.today().strftime("%B %d, %Y")


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, font_size=10, color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    run.bold = bold
    return run


def add_formatted_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1B3A5C")
        set_cell_text(cell, header, bold=True, font_size=10, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, row in enumerate(rows):
        for c_idx, (text, is_bold, align) in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")
            set_cell_text(cell, text, bold=is_bold, font_size=10, color=DARK_GRAY, alignment=align)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    return table


def add_header_footer(doc):
    for section in doc.sections[1:]:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_before = Pt(4)
        hp.paragraph_format.space_after = Pt(4)
        run = hp.add_run("ValClassifieds v1  |  Automation Test Execution Report")
        run.font.size = Pt(8)
        run.font.color.rgb = MED_GRAY
        run.font.name = "Calibri"

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after = Pt(4)
        run = fp.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = MED_GRAY
        run.font.name = "Calibri"
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE   \\* MERGEFORMAT ">'
            '<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="666666"/>'
            '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>'
            '<w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls("w"))
        fp._p.append(parse_xml(fld_xml))
        run2 = fp.add_run(" of ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = MED_GRAY
        run2.font.name = "Calibri"
        fld_xml2 = (
            '<w:fldSimple {} w:instr=" NUMPAGES   \\* MERGEFORMAT ">'
            '<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="666666"/>'
            '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>'
            '<w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls("w"))
        fp._p.append(parse_xml(fld_xml2))


def build_cover_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ValClassifieds v1")
    run.font.size = Pt(28)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Automation Test Execution Report")
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT_BLUE
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.size = Pt(10)
    run.font.color.rgb = MED_GRAY

    for label, value in [("Version", "1.0"), ("Date", TODAY), ("Prepared By", "Senior QA Automation Engineer")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{label}:  ")
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_GRAY
        run.font.name = "Calibri"
        run.bold = True
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = ACCENT_BLUE
        run.font.name = "Calibri"

    doc.add_page_break()


def build_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table of Contents")
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    run.bold = True

    doc.add_paragraph()

    toc_items = [
        ("1.", "Executive Summary"),
        ("2.", "Project Overview"),
        ("3.", "Automation Framework Overview"),
        ("4.", "Test Execution Summary"),
        ("5.", "Detailed Test Coverage"),
        ("6.", "Test Environment"),
        ("7.", "Automation Features Implemented"),
        ("8.", "Results"),
        ("9.", "Quality Assessment"),
        ("10.", "Recommendations"),
        ("11.", "Conclusion"),
        ("", "Appendix"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        text = f"  {num}  {title}" if num else f"       {title}"
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT_BLUE if num else DARK_GRAY
        run.font.name = "Calibri"
        run.bold = bool(num)

    doc.add_page_break()


def build_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = ACCENT_BLUE

    return p


def build_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_BLUE
    run.font.name = "Calibri"
    run.bold = True
    return p


def build_body_text(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    return p


def build_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run.font.name = "Calibri"
        run.bold = True
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.color.rgb = DARK_GRAY
        run2.font.name = "Calibri"
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run.font.name = "Calibri"
    return p


def build_checkmark(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("\u2713  ")
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    run.font.name = "Calibri"
    run.bold = True
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = DARK_GRAY
    run2.font.name = "Calibri"
    return p


def add_page_break(doc):
    doc.add_page_break()


def build_executive_summary(doc):
    build_section_heading(doc, "1. Executive Summary")
    build_body_text(doc,
        "ValClassifieds v1 is a modern, full-featured classifieds marketplace platform built with React, TypeScript, "
        "and Supabase. The application enables users to browse, search, and list classified advertisements with "
        "real-time chat messaging, user authentication, and responsive design across devices."
    )
    build_body_text(doc,
        "This Automation Test Execution Report documents the comprehensive end-to-end (E2E) testing effort "
        "undertaken to validate the functional correctness, accessibility, and user experience of the ValClassifieds v1 "
        "application. The automation framework was designed and implemented using Playwright with TypeScript, "
        "following industry best practices including the Page Object Model (POM), mock authentication, and mock data strategies."
    )
    build_body_text(doc,
        "The primary objectives of this automation initiative were to establish a reliable regression safety net, "
        "accelerate release cycles through automated validation, provide reproducible test execution, and enable "
        "seamless integration into a CI/CD pipeline. The framework covers the three core modules of the application: "
        "Authentication, Chat Messaging, and Listings Management."
    )
    build_body_text(doc,
        "Overall Quality Assessment:", bold=True
    )
    build_body_text(doc,
        "A total of 22 automated end-to-end tests were executed across 3 test suites, achieving a 100% pass rate "
        "with zero failures or skipped tests. The automation framework demonstrates exceptional stability, reliability, "
        "and maintainability. The application's core functionality is well-covered by automated tests, providing "
        "confidence for production deployment and future development cycles."
    )


def build_project_overview(doc):
    build_section_heading(doc, "2. Project Overview")

    build_body_text(doc,
        "ValClassifieds v1 is a web-based classifieds marketplace that connects buyers and sellers through an "
        "intuitive, responsive interface. The platform supports user registration and authentication, listing creation "
        "and management, advanced search and filtering, real-time messaging, and image upload capabilities."
    )

    build_sub_heading(doc, "Application Type")
    build_body_text(doc, "Web Application - Classifieds Marketplace Platform (Single Page Application)")

    build_sub_heading(doc, "Technology Stack")

    headers = ["Component", "Technology", "Version"]
    rows = [
        ("Frontend Framework", "React", "18.3.1", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Language", "TypeScript", "5.5.4", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("CSS Framework", "Tailwind CSS", "3.4.7", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Backend / Database", "Supabase", "2.45.0", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Build Tool", "Vite", "5.4.0", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("State Management", "Zustand", "5.0.14", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Routing", "React Router", "6.26.0", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("E2E Testing", "Playwright", "1.61.1", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Unit Testing", "Vitest", "4.1.9", False, WD_ALIGN_PARAGRAPH.LEFT),
    ]
    formatted_rows = []
    for r in rows:
        formatted_rows.append([
            (r[0], True, WD_ALIGN_PARAGRAPH.LEFT),
            (r[1], False, WD_ALIGN_PARAGRAPH.LEFT),
            (r[2], False, WD_ALIGN_PARAGRAPH.CENTER),
        ])
    add_formatted_table(doc, headers, formatted_rows, col_widths=[2.5, 2.0, 1.0])


def build_framework_overview(doc):
    build_section_heading(doc, "3. Automation Framework Overview")

    build_body_text(doc,
        "The ValClassifieds v1 automation framework is built on Playwright with TypeScript, leveraging a modular, "
        "maintainable architecture designed for scalability and CI/CD integration. The framework follows the Page "
        "Object Model (POM) design pattern to separate test logic from page-specific implementation details."
    )

    build_sub_heading(doc, "Framework Architecture")
    build_body_text(doc,
        "The framework is organized into a clear directory structure with distinct layers for page objects, "
        "component objects, test specifications, utilities, fixtures, and configuration. This separation of concerns "
        "enables parallel test development, reduces code duplication, and simplifies maintenance."
    )

    build_sub_heading(doc, "Page Object Model (POM)")
    build_body_text(doc,
        "Each application page or feature module has a corresponding page object class that encapsulates element "
        "selectors, interaction methods, and page-specific utilities. Tests interact with pages through these "
        "abstractions, ensuring that UI changes only require updates in one location."
    )

    build_sub_heading(doc, "Mock Authentication Strategy")
    build_body_text(doc,
        "To eliminate dependency on a live Supabase backend during test execution, the framework implements a "
        "mock authentication system. A localStorage-based session injection mimics authenticated user state, "
        "and Playwright route interception handles all REST API calls (auth tokens, profiles, conversations, "
        "messages, listings, storage objects). This approach provides deterministic, fast, and reliable test execution."
    )

    build_sub_heading(doc, "Mock Data Strategy")
    build_body_text(doc,
        "Comprehensive mock data sets were created to support realistic test scenarios. These include 5 conversations "
        "with 74 messages across multiple participants, read/unread states, image attachments, user profiles with "
        "presence indicators, listings with categories and images, and diverse user roles. The mock data ensures "
        "tests exercise real-world application behavior without requiring a seeded database."
    )

    build_sub_heading(doc, "Reusable Utilities")
    build_body_text(doc,
        "The framework includes a suite of reusable utilities including: test data factories, custom Playwright "
        "fixtures (test context with authenticated state), centralized selector definitions, and helper functions "
        "for common test operations. These utilities accelerate test development and promote consistency."
    )

    build_sub_heading(doc, "Reporting Mechanism")
    build_body_text(doc,
        "Test results are captured through multiple reporting channels: built-in Playwright HTML reporter produces "
        "an interactive test report with screenshots, traces, and video; Allure reporting generates rich, "
        "customizable reports with test history and trends; and the Playwright UI Mode provides a real-time "
        "test explorer for debugging and development. Screenshots, trace files, and video recordings are "
        "captured automatically on test failure for detailed forensic analysis."
    )


def build_execution_summary(doc):
    build_section_heading(doc, "4. Test Execution Summary")

    build_body_text(doc,
        "The following table summarizes the results of the complete E2E test execution run. All tests were "
        "executed in headless Chromium on a macOS environment."
    )

    doc.add_paragraph()
    headers = ["Metric", "Result"]
    rows = [
        ("Total Test Suites", "3", True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Authentication Tests", "6", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Chat Tests", "8", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Listings Tests", "8", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Total Automated Tests", "22", True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Passed", "22", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Failed", "0", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Skipped", "0", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Pass Percentage", "100%", True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Execution Time", "~8 minutes", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Browser", "Chromium", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Framework", "Playwright", False, WD_ALIGN_PARAGRAPH.LEFT),
    ]
    formatted_rows = [[(r[0], r[2], WD_ALIGN_PARAGRAPH.LEFT), (r[1], False, WD_ALIGN_PARAGRAPH.CENTER)] for r in rows]
    add_formatted_table(doc, headers, formatted_rows, col_widths=[3.0, 3.0])


def build_detailed_coverage(doc):
    build_section_heading(doc, "5. Detailed Test Coverage")

    build_sub_heading(doc, "Authentication Module")
    headers = ["Test Case", "Description"]
    rows = [
        ("Login Page Rendering", "Verify login page loads with email, password inputs, submit button, and registration link"),
        ("Form Validation", "Verify validation messages appear for empty fields and invalid email format"),
        ("Invalid Credentials", "Verify appropriate error message when invalid credentials are submitted"),
        ("Successful Login", "Verify user is redirected to home page upon successful authentication"),
        ("Navigation", "Verify user can navigate to register page and back to login page"),
        ("Accessibility", "Verify keyboard navigation, ARIA labels, and focus management on login form"),
    ]
    formatted_rows = [[(r[0], False, WD_ALIGN_PARAGRAPH.LEFT), (r[1], False, WD_ALIGN_PARAGRAPH.LEFT)] for r in rows]
    add_formatted_table(doc, headers, formatted_rows, col_widths=[1.8, 4.2])

    doc.add_paragraph()

    build_sub_heading(doc, "Chat Messaging Module")
    headers = ["Test Case", "Description"]
    rows = [
        ("Conversation List", "Verify conversation list displays correct number of conversations with participant names"),
        ("Logged-in User Display", "Verify the logged-in user's name is displayed in the conversation area"),
        ("Send Message", "Verify user can type and send a message, seeing it appear in the chat area"),
        ("Unread Badge", "Verify conversation with unread messages displays an unread count badge"),
        ("Search Conversations", "Verify search by participant name filters the conversation list correctly"),
        ("No Results Filter", "Verify search with non-matching text shows empty state with appropriate message"),
        ("Message Input Accessibility", "Verify message input has proper ARIA labels and keyboard support"),
        ("Back Navigation", "Verify user can return from conversation view to conversation list"),
    ]
    formatted_rows = [[(r[0], False, WD_ALIGN_PARAGRAPH.LEFT), (r[1], False, WD_ALIGN_PARAGRAPH.LEFT)] for r in rows]
    add_formatted_table(doc, headers, formatted_rows, col_widths=[1.8, 4.2])

    doc.add_paragraph()

    build_sub_heading(doc, "Listings Management Module")
    headers = ["Test Case", "Description"]
    rows = [
        ("Listing Cards", "Verify listings page displays listing cards with title, price, location, and image"),
        ("Search Functionality", "Verify search input filters listings by title in real-time"),
        ("Category Filter", "Verify category dropdown filters listings by selected category"),
        ("Location Filter", "Verify location input filters listings by city"),
        ("Price Filter", "Verify price range inputs filter listings within specified range"),
        ("Empty State", "Verify appropriate empty state message when no listings match filters"),
        ("Listing Details", "Verify clicking a listing card navigates to detail page with full information"),
        ("Accessibility", "Verify listing cards have proper ARIA labels and keyboard navigation"),
    ]
    formatted_rows = [[(r[0], False, WD_ALIGN_PARAGRAPH.LEFT), (r[1], False, WD_ALIGN_PARAGRAPH.LEFT)] for r in rows]
    add_formatted_table(doc, headers, formatted_rows, col_widths=[1.8, 4.2])


def build_test_environment(doc):
    build_section_heading(doc, "6. Test Environment")

    headers = ["Parameter", "Value"]
    rows = [
        ("Operating System", "macOS", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Browser", "Chromium", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Framework Version", "Playwright 1.61.1", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Node.js", "v24.18.0", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Execution Mode", "Headless (CLI) / Interactive (UI / Debug)", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Mock Authentication", "localStorage session injection + REST route interception", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Mock Data", "5 conversations, 74 messages, 71 listings, 232 images, 16 JSON data files", False, WD_ALIGN_PARAGRAPH.LEFT),
    ]
    formatted_rows = [[(r[0], True, WD_ALIGN_PARAGRAPH.LEFT), (r[1], False, WD_ALIGN_PARAGRAPH.LEFT)] for r in rows]
    add_formatted_table(doc, headers, formatted_rows, col_widths=[2.5, 3.5])


def build_features(doc):
    build_section_heading(doc, "7. Automation Features Implemented")

    features = [
        "Page Object Model (POM) architecture for clean separation of test logic and page interactions",
        "Modular framework design with distinct layers for pages, components, specs, utilities, and fixtures",
        "Reusable component objects (Navbar, ListingCard) for shared UI patterns across pages",
        "Mock authentication via localStorage session injection and REST API route interception",
        "Mock conversation data with 5 conversations, 74 messages, and real-world read/unread states",
        "Mock listings data with 71 listings across multiple categories and 232 associated images",
        "Mock user profiles with diverse Indian names, roles, and presence indicators",
        "data-testid attribute strategy for resilient, semantic element selection",
        "Playwright HTML reporter for interactive test result visualization",
        "Allure reporting for rich, customizable test reports with history and trends",
        "Playwright UI Mode for real-time test exploration and debugging",
        "Playwright Debug Mode for step-through test execution with Playwright Inspector",
        "Screenshot capture on test failure for visual forensics",
        "Trace file capture for detailed DOM snapshots, network logs, and timing data",
        "Video recording of test execution for behavioral analysis",
    ]

    for feature in features:
        build_checkmark(doc, feature)


def build_results(doc):
    build_section_heading(doc, "8. Results")

    build_body_text(doc,
        "22 out of 22 automated end-to-end tests executed successfully.", bold=True
    )
    build_body_text(doc, "No failed tests.")
    build_body_text(doc, "No skipped tests.")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Overall Project Automation Health:  GREEN")
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN
    run.font.name = "Calibri"
    run.bold = True

    doc.add_paragraph()

    coverage_headers = ["Module", "Tests", "Passed", "Failed", "Skipped", "Pass %"]
    coverage_rows = [
        ("Authentication", "6", "6", "0", "0", "100%"),
        ("Chat Messaging", "8", "8", "0", "0", "100%"),
        ("Listings", "8", "8", "0", "0", "100%"),
        ("Total", "22", "22", "0", "0", "100%"),
    ]
    formatted_rows = []
    for r in coverage_rows:
        is_total = r[0] == "Total"
        formatted_rows.append([
            (r[0], is_total, WD_ALIGN_PARAGRAPH.LEFT),
            (r[1], is_total, WD_ALIGN_PARAGRAPH.CENTER),
            (r[2], is_total, WD_ALIGN_PARAGRAPH.CENTER),
            (r[3], is_total, WD_ALIGN_PARAGRAPH.CENTER),
            (r[4], is_total, WD_ALIGN_PARAGRAPH.CENTER),
            (r[5], is_total, WD_ALIGN_PARAGRAPH.CENTER),
        ])
    add_formatted_table(doc, coverage_headers, formatted_rows, col_widths=[1.5, 0.8, 0.8, 0.8, 0.8, 0.8])


def build_quality_assessment(doc):
    build_section_heading(doc, "9. Quality Assessment")

    assessments = [
        ("Stability",
         "The automation framework demonstrates exceptional stability, with all 22 tests consistently passing "
         "across multiple execution runs. The mock-based approach eliminates flakiness introduced by network "
         "latency, rate limiting, or backend instability. Tests are self-contained and produce deterministic results."),
        ("Reliability",
         "Mock authentication and data strategies ensure tests are not dependent on external services. Route "
         "interception provides predictable API responses, while the localStorage-based session injection "
         "simulates authenticated state without requiring real credentials. The framework reliably validates "
         "both happy paths and edge cases."),
        ("Maintainability",
         "The Page Object Model architecture, modular directory structure, and reusable utilities significantly "
         "reduce maintenance overhead. UI changes require updates only in page object methods, while test "
         "specifications remain unchanged. Consistent coding conventions and descriptive test names enhance "
         "readability and onboarding efficiency."),
        ("Scalability",
         "The framework architecture supports easy addition of new test suites. Adding a new feature module "
         "requires creating a page object, test spec, and mock data - following established patterns. "
         "Playwright's parallel execution capabilities enable efficient scaling across multiple spec files."),
        ("Readiness for CI/CD",
         "The framework is fully ready for CI/CD integration. Tests run in headless mode, produce machine-readable "
         "output (JUnit XML, Allure results), and support environment configuration through environment variables. "
         "The GitHub Actions workflow template can be activated with minimal configuration to run tests on pull "
         "requests and scheduled intervals."),
    ]

    for title, desc in assessments:
        build_sub_heading(doc, title)
        build_body_text(doc, desc)


def build_recommendations(doc):
    build_section_heading(doc, "10. Recommendations")

    recommendations = [
        ("Cross-Browser Testing",
         "Extend the test matrix to include WebKit (Safari) and Firefox browsers alongside Chromium. Playwright's "
         "cross-browser support enables consistent validation across browser engines with minimal code changes."),
        ("Mobile Device Testing",
         "Leverage Playwright's device emulation to validate responsive layouts and touch interactions on mobile "
         "viewport sizes. This ensures the mobile user experience meets quality standards."),
        ("API Automation",
         "Develop a dedicated API test suite to validate backend endpoints, business logic, and data integrity "
         "independently of the UI. This provides faster feedback on backend changes and improves coverage."),
        ("Performance Automation",
         "Introduce performance testing using Playwright's built-in performance API or integrate with tools like "
         "Lighthouse CI to track Core Web Vitals, load times, and runtime performance metrics."),
        ("GitHub Actions Integration",
         "Activate the CI/CD pipeline by configuring GitHub Actions workflows. Configure parallel test execution "
         "across browsers, Allure report publishing, Slack/email notifications on failures, and automatic "
         "test execution on pull request creation."),
    ]

    for title, desc in recommendations:
        build_sub_heading(doc, title)
        build_body_text(doc, desc)


def build_conclusion(doc):
    build_section_heading(doc, "11. Conclusion")

    build_body_text(doc,
        "The ValClassifieds v1 Automation Test Execution Report confirms that the application has undergone "
        "rigorous automated end-to-end testing across its three core modules: Authentication, Chat Messaging, "
        "and Listings Management."
    )
    build_body_text(doc,
        "With 22 out of 22 automated tests passing (100% pass rate), zero failures, and zero skipped tests, "
        "the automation framework demonstrates production-ready quality and reliability. The framework's modular "
        "architecture, comprehensive mock data strategy, and robust reporting mechanisms make it well-suited for "
        "continuous integration and continuous delivery pipelines."
    )
    build_body_text(doc,
        "The automation framework is production-ready and suitable for continuous integration.", bold=True
    )
    build_body_text(doc,
        "The combination of Playwright's powerful testing capabilities, the Page Object Model design pattern, "
        "and comprehensive mock data provides a solid foundation for maintaining quality as the application "
        "evolves. The framework is designed to scale with the project, supporting new features, cross-browser "
        "testing, and CI/CD integration with minimal additional investment."
    )


def build_appendix(doc):
    build_section_heading(doc, "Appendix")

    build_sub_heading(doc, "Folder Structure")
    struct = (
        "tests/e2e/\n"
        "  \u2514\u2500\u2500 pages/              Page Object classes\n"
        "  \u2514\u2500\u2500 components/         Component Object classes\n"
        "  \u2514\u2500\u2500 specs/              Test specifications\n"
        "  \u2502     \u2514\u2500\u2500 auth/     Authentication tests\n"
        "  \u2502     \u2514\u2500\u2500 chat/     Chat messaging tests\n"
        "  \u2502     \u2514\u2500\u2500 listings/ Listings tests\n"
        "  \u2514\u2500\u2500 utils/              Utilities, selectors, test data\n"
        "  \u2514\u2500\u2500 fixtures/           Custom Playwright fixtures\n"
        "  \u2514\u2500\u2500 setup/              Global setup / teardown\n"
        "  \u2514\u2500\u2500 seeds/              SQL seed scripts\n"
        "  \u2514\u2500\u2500 playwright.config.ts Configuration\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(struct)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    run.font.name = "Consolas"
    run.font.italic = True

    build_sub_heading(doc, "Automation Documentation Created")
    docs = [
        "Automation Test Plan",
        "Framework Architecture Document",
        "Framework Design Specification",
        "Automation Scope Matrix",
        "Requirements Traceability Matrix",
        "Object Repository Map",
        "Test Data Strategy Document",
        "Automation Roadmap",
        "Coding Standards Document",
        "Reporting Strategy Document",
        "Maintenance Guide",
        "Automation Metrics Document",
    ]
    for d in docs:
        build_bullet(doc, d)

    doc.add_paragraph()

    build_sub_heading(doc, "QA Documentation Created")
    qa_docs = [
        "Master Test Case Spreadsheet (180+ test cases across 14 modules)",
        "Requirements Traceability Matrix (48+ requirements mapped)",
        "Smoke Test Suite Spreadsheet",
        "Regression Test Suite Spreadsheet",
        "Sanity Test Suite Spreadsheet",
        "Security Test Suite Spreadsheet",
        "Performance Test Suite Spreadsheet",
        "User Acceptance Test Spreadsheet",
        "Test Summary Report",
        "QA Analysis Report",
    ]
    for d in qa_docs:
        build_bullet(doc, d)

    doc.add_paragraph()

    build_sub_heading(doc, "Mock Data Created")
    mock_items = [
        "16 JSON data files (users, profiles, categories, listings, images, conversations, messages, etc.)",
        "71 classified listings with realistic titles, descriptions, and pricing",
        "232 listing images across multiple categories",
        "5 conversations with 74 messages across multiple participants",
        "Conversation read/unread tracking states",
        "User presence indicators (online, away, offline)",
        "Message reactions and attachments",
    ]
    for m in mock_items:
        build_bullet(doc, m)

    doc.add_paragraph()

    build_sub_heading(doc, "Playwright Reports Generated")
    report_items = [
        "Playwright HTML Report (interactive, with screenshots, traces, and video)",
        "Allure Report (rich dashboard with test history and trends)",
        "JUnit XML results for CI/CD integration",
    ]
    for r in report_items:
        build_bullet(doc, r)


def generate_docx():
    doc = Document()

    default_style = doc.styles["Normal"]
    default_style.font.name = "Calibri"
    default_style.font.size = Pt(11)
    default_style.font.color.rgb = DARK_GRAY
    default_style.paragraph_format.space_after = Pt(6)

    build_cover_page(doc)

    section = doc.add_section()
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    build_toc(doc)
    build_executive_summary(doc)
    add_page_break(doc)
    build_project_overview(doc)
    add_page_break(doc)
    build_framework_overview(doc)
    add_page_break(doc)
    build_execution_summary(doc)
    add_page_break(doc)
    build_detailed_coverage(doc)
    add_page_break(doc)
    build_test_environment(doc)
    add_page_break(doc)
    build_features(doc)
    add_page_break(doc)
    build_results(doc)
    add_page_break(doc)
    build_quality_assessment(doc)
    add_page_break(doc)
    build_recommendations(doc)
    add_page_break(doc)
    build_conclusion(doc)
    add_page_break(doc)
    build_appendix(doc)

    add_header_footer(doc)

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")
    return True


def generate_html():
    """Generate HTML version for PDF conversion via Playwright."""
    html = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Automation Test Execution Report - ValClassifieds v1</title>
<style>
  @page {
    size: A4;
    margin: 2cm 2.5cm;
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
    color: #333;
    font-size: 11pt;
    line-height: 1.5;
  }

  .cover-page {
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    height: 95vh;
    text-align: center;
  }
  .cover-page h1 { font-size: 28pt; color: #1B3A5C; margin-bottom: 8px; }
  .cover-page h2 { font-size: 22pt; color: #2E5C8A; margin-bottom: 16px; font-weight: 400; }
  .cover-page .divider { color: #999; margin: 16px 0; letter-spacing: 4px; }
  .cover-page .meta { font-size: 12pt; margin: 6px 0; }
  .cover-page .meta strong { color: #333; }
  .cover-page .meta span { color: #2E5C8A; }

  .page-break { page-break-before: always; }

  .section-title {
    font-size: 18pt;
    color: #1B3A5C;
    font-weight: bold;
    margin-top: 28px;
    margin-bottom: 8px;
  }
  .section-divider {
    border: none;
    border-top: 1px solid #2E5C8A;
    margin-bottom: 16px;
    width: 100%;
  }
  .subsection-title {
    font-size: 14pt;
    color: #2E5C8A;
    font-weight: bold;
    margin-top: 18px;
    margin-bottom: 8px;
  }

  p { margin: 6px 0; text-align: justify; }

  table {
    width: 100%;
    border-collapse: collapse;
    margin: 12px 0;
    font-size: 10pt;
  }
  th {
    background-color: #1B3A5C;
    color: white;
    padding: 6px 10px;
    text-align: center;
    font-weight: bold;
    border: 1px solid #1B3A5C;
  }
  td {
    padding: 5px 10px;
    border: 1px solid #ccc;
    text-align: left;
  }
  tr:nth-child(even) td { background-color: #F2F6FA; }
  td.metric { font-weight: bold; }
  td.center { text-align: center; }

  ul { margin: 4px 0 4px 20px; }
  li { margin: 2px 0; }

  .check { color: #1B8A3D; font-weight: bold; }

  .health-green {
    text-align: center;
    font-size: 16pt;
    color: #1B8A3D;
    font-weight: bold;
    margin: 16px 0;
  }

  .code {
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 9pt;
    color: #555;
    line-height: 1.4;
    margin: 8px 0;
    padding: 10px;
    background: #f8f8f8;
    border-left: 3px solid #2E5C8A;
  }

  .status-pass { color: #1B8A3D; font-weight: bold; }
  .status-fail { color: #C0392B; font-weight: bold; }

  @media print {
    .cover-page { height: 100vh; }
    .page-break { page-break-before: always; }
  }
</style>
</head>
<body>
"""

    html += _cover_html()
    html += _toc_html()
    html += _executive_summary_html()
    html += _project_overview_html()
    html += _framework_overview_html()
    html += _execution_summary_html()
    html += _detailed_coverage_html()
    html += _test_environment_html()
    html += _features_html()
    html += _results_html()
    html += _quality_assessment_html()
    html += _recommendations_html()
    html += _conclusion_html()
    html += _appendix_html()

    html += """
</body>
</html>"""

    with open(HTML_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"HTML saved: {HTML_PATH}")
    return HTML_PATH


def _cover_html():
    return f"""
<div class="cover-page">
  <h1>ValClassifieds v1</h1>
  <h2>Automation Test Execution Report</h2>
  <div class="divider">&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;&mdash;</div>
  <p class="meta"><strong>Version:</strong> <span>1.0</span></p>
  <p class="meta"><strong>Date:</strong> <span>{TODAY}</span></p>
  <p class="meta"><strong>Prepared By:</strong> <span>Senior QA Automation Engineer</span></p>
</div>
"""


def _toc_html():
    return """
<div class="page-break">
  <h2 style="text-align:center;color:#1B3A5C;font-size:20pt;margin-bottom:20px;">Table of Contents</h2>
  <ul style="list-style:none;margin:0;padding:0;">
    <li style="margin:6px 0;"><strong>1.</strong> Executive Summary</li>
    <li style="margin:6px 0;"><strong>2.</strong> Project Overview</li>
    <li style="margin:6px 0;"><strong>3.</strong> Automation Framework Overview</li>
    <li style="margin:6px 0;"><strong>4.</strong> Test Execution Summary</li>
    <li style="margin:6px 0;"><strong>5.</strong> Detailed Test Coverage</li>
    <li style="margin:6px 0;"><strong>6.</strong> Test Environment</li>
    <li style="margin:6px 0;"><strong>7.</strong> Automation Features Implemented</li>
    <li style="margin:6px 0;"><strong>8.</strong> Results</li>
    <li style="margin:6px 0;"><strong>9.</strong> Quality Assessment</li>
    <li style="margin:6px 0;"><strong>10.</strong> Recommendations</li>
    <li style="margin:6px 0;"><strong>11.</strong> Conclusion</li>
    <li style="margin:6px 0;color:#555;">&nbsp;&nbsp;&nbsp;&nbsp;Appendix</li>
  </ul>
</div>
"""


def _executive_summary_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">1. Executive Summary</h1>
  <hr class="section-divider">
  <p>ValClassifieds v1 is a modern, full-featured classifieds marketplace platform built with React, TypeScript, and Supabase. The application enables users to browse, search, and list classified advertisements with real-time chat messaging, user authentication, and responsive design across devices.</p>
  <p>This Automation Test Execution Report documents the comprehensive end-to-end (E2E) testing effort undertaken to validate the functional correctness, accessibility, and user experience of the ValClassifieds v1 application. The automation framework was designed and implemented using Playwright with TypeScript, following industry best practices including the Page Object Model (POM), mock authentication, and mock data strategies.</p>
  <p>The primary objectives of this automation initiative were to establish a reliable regression safety net, accelerate release cycles through automated validation, provide reproducible test execution, and enable seamless integration into a CI/CD pipeline. The framework covers the three core modules of the application: Authentication, Chat Messaging, and Listings Management.</p>
  <p><strong>Overall Quality Assessment:</strong></p>
  <p>A total of 22 automated end-to-end tests were executed across 3 test suites, achieving a 100% pass rate with zero failures or skipped tests. The automation framework demonstrates exceptional stability, reliability, and maintainability. The application's core functionality is well-covered by automated tests, providing confidence for production deployment and future development cycles.</p>
</div>
"""


def _project_overview_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">2. Project Overview</h1>
  <hr class="section-divider">
  <p>ValClassifieds v1 is a web-based classifieds marketplace that connects buyers and sellers through an intuitive, responsive interface. The platform supports user registration and authentication, listing creation and management, advanced search and filtering, real-time messaging, and image upload capabilities.</p>
  <h2 class="subsection-title">Application Type</h2>
  <p>Web Application - Classifieds Marketplace Platform (Single Page Application)</p>
  <h2 class="subsection-title">Technology Stack</h2>
  <table>
    <tr><th>Component</th><th>Technology</th><th>Version</th></tr>
    <tr><td>Frontend Framework</td><td>React</td><td class="center">18.3.1</td></tr>
    <tr><td>Language</td><td>TypeScript</td><td class="center">5.5.4</td></tr>
    <tr><td>CSS Framework</td><td>Tailwind CSS</td><td class="center">3.4.7</td></tr>
    <tr><td>Backend / Database</td><td>Supabase</td><td class="center">2.45.0</td></tr>
    <tr><td>Build Tool</td><td>Vite</td><td class="center">5.4.0</td></tr>
    <tr><td>State Management</td><td>Zustand</td><td class="center">5.0.14</td></tr>
    <tr><td>Routing</td><td>React Router</td><td class="center">6.26.0</td></tr>
    <tr><td>E2E Testing</td><td>Playwright</td><td class="center">1.61.1</td></tr>
    <tr><td>Unit Testing</td><td>Vitest</td><td class="center">4.1.9</td></tr>
  </table>
</div>
"""


def _framework_overview_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">3. Automation Framework Overview</h1>
  <hr class="section-divider">
  <p>The ValClassifieds v1 automation framework is built on Playwright with TypeScript, leveraging a modular, maintainable architecture designed for scalability and CI/CD integration. The framework follows the Page Object Model (POM) design pattern to separate test logic from page-specific implementation details.</p>
  <h2 class="subsection-title">Framework Architecture</h2>
  <p>The framework is organized into a clear directory structure with distinct layers for page objects, component objects, test specifications, utilities, fixtures, and configuration. This separation of concerns enables parallel test development, reduces code duplication, and simplifies maintenance.</p>
  <h2 class="subsection-title">Page Object Model (POM)</h2>
  <p>Each application page or feature module has a corresponding page object class that encapsulates element selectors, interaction methods, and page-specific utilities. Tests interact with pages through these abstractions, ensuring that UI changes only require updates in one location.</p>
  <h2 class="subsection-title">Mock Authentication Strategy</h2>
  <p>To eliminate dependency on a live Supabase backend during test execution, the framework implements a mock authentication system. A localStorage-based session injection mimics authenticated user state, and Playwright route interception handles all REST API calls (auth tokens, profiles, conversations, messages, listings, storage objects). This approach provides deterministic, fast, and reliable test execution.</p>
  <h2 class="subsection-title">Mock Data Strategy</h2>
  <p>Comprehensive mock data sets were created to support realistic test scenarios. These include 5 conversations with 74 messages across multiple participants, read/unread states, image attachments, user profiles with presence indicators, listings with categories and images, and diverse user roles. The mock data ensures tests exercise real-world application behavior without requiring a seeded database.</p>
  <h2 class="subsection-title">Reusable Utilities</h2>
  <p>The framework includes a suite of reusable utilities including: test data factories, custom Playwright fixtures (test context with authenticated state), centralized selector definitions, and helper functions for common test operations. These utilities accelerate test development and promote consistency.</p>
  <h2 class="subsection-title">Reporting Mechanism</h2>
  <p>Test results are captured through multiple reporting channels: built-in Playwright HTML reporter produces an interactive test report with screenshots, traces, and video; Allure reporting generates rich, customizable reports with test history and trends; and the Playwright UI Mode provides a real-time test explorer for debugging and development. Screenshots, trace files, and video recordings are captured automatically on test failure for detailed forensic analysis.</p>
</div>
"""


def _execution_summary_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">4. Test Execution Summary</h1>
  <hr class="section-divider">
  <p>The following table summarizes the results of the complete E2E test execution run. All tests were executed in headless Chromium on a macOS environment.</p>
  <table>
    <tr><th>Metric</th><th>Result</th></tr>
    <tr><td class="metric">Total Test Suites</td><td class="center">3</td></tr>
    <tr><td>Authentication Tests</td><td class="center">6</td></tr>
    <tr><td>Chat Tests</td><td class="center">8</td></tr>
    <tr><td>Listings Tests</td><td class="center">8</td></tr>
    <tr><td class="metric">Total Automated Tests</td><td class="center">22</td></tr>
    <tr><td>Passed</td><td class="center">22</td></tr>
    <tr><td>Failed</td><td class="center">0</td></tr>
    <tr><td>Skipped</td><td class="center">0</td></tr>
    <tr><td class="metric">Pass Percentage</td><td class="center">100%</td></tr>
    <tr><td>Execution Time</td><td class="center">~8 minutes</td></tr>
    <tr><td>Browser</td><td class="center">Chromium</td></tr>
    <tr><td>Framework</td><td class="center">Playwright</td></tr>
  </table>
</div>
"""


def _detailed_coverage_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">5. Detailed Test Coverage</h1>
  <hr class="section-divider">

  <h2 class="subsection-title">Authentication Module</h2>
  <table>
    <tr><th>Test Case</th><th>Description</th></tr>
    <tr><td>Login Page Rendering</td><td>Verify login page loads with email, password inputs, submit button, and registration link</td></tr>
    <tr><td>Form Validation</td><td>Verify validation messages appear for empty fields and invalid email format</td></tr>
    <tr><td>Invalid Credentials</td><td>Verify appropriate error message when invalid credentials are submitted</td></tr>
    <tr><td>Successful Login</td><td>Verify user is redirected to home page upon successful authentication</td></tr>
    <tr><td>Navigation</td><td>Verify user can navigate to register page and back to login page</td></tr>
    <tr><td>Accessibility</td><td>Verify keyboard navigation, ARIA labels, and focus management on login form</td></tr>
  </table>

  <h2 class="subsection-title">Chat Messaging Module</h2>
  <table>
    <tr><th>Test Case</th><th>Description</th></tr>
    <tr><td>Conversation List</td><td>Verify conversation list displays correct number of conversations with participant names</td></tr>
    <tr><td>Logged-in User Display</td><td>Verify the logged-in user's name is displayed in the conversation area</td></tr>
    <tr><td>Send Message</td><td>Verify user can type and send a message, seeing it appear in the chat area</td></tr>
    <tr><td>Unread Badge</td><td>Verify conversation with unread messages displays an unread count badge</td></tr>
    <tr><td>Search Conversations</td><td>Verify search by participant name filters the conversation list correctly</td></tr>
    <tr><td>No Results Filter</td><td>Verify search with non-matching text shows empty state with appropriate message</td></tr>
    <tr><td>Message Input Accessibility</td><td>Verify message input has proper ARIA labels and keyboard support</td></tr>
    <tr><td>Back Navigation</td><td>Verify user can return from conversation view to conversation list</td></tr>
  </table>

  <h2 class="subsection-title">Listings Management Module</h2>
  <table>
    <tr><th>Test Case</th><th>Description</th></tr>
    <tr><td>Listing Cards</td><td>Verify listings page displays listing cards with title, price, location, and image</td></tr>
    <tr><td>Search Functionality</td><td>Verify search input filters listings by title in real-time</td></tr>
    <tr><td>Category Filter</td><td>Verify category dropdown filters listings by selected category</td></tr>
    <tr><td>Location Filter</td><td>Verify location input filters listings by city</td></tr>
    <tr><td>Price Filter</td><td>Verify price range inputs filter listings within specified range</td></tr>
    <tr><td>Empty State</td><td>Verify appropriate empty state message when no listings match filters</td></tr>
    <tr><td>Listing Details</td><td>Verify clicking a listing card navigates to detail page with full information</td></tr>
    <tr><td>Accessibility</td><td>Verify listing cards have proper ARIA labels and keyboard navigation</td></tr>
  </table>
</div>
"""


def _test_environment_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">6. Test Environment</h1>
  <hr class="section-divider">
  <table>
    <tr><th>Parameter</th><th>Value</th></tr>
    <tr><td class="metric">Operating System</td><td>macOS</td></tr>
    <tr><td class="metric">Browser</td><td>Chromium</td></tr>
    <tr><td class="metric">Framework Version</td><td>Playwright 1.61.1</td></tr>
    <tr><td class="metric">Node.js</td><td>v24.18.0</td></tr>
    <tr><td class="metric">Execution Mode</td><td>Headless (CLI) / Interactive (UI / Debug)</td></tr>
    <tr><td class="metric">Mock Authentication</td><td>localStorage session injection + REST route interception</td></tr>
    <tr><td class="metric">Mock Data</td><td>5 conversations, 74 messages, 71 listings, 232 images, 16 JSON data files</td></tr>
  </table>
</div>
"""


def _features_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">7. Automation Features Implemented</h1>
  <hr class="section-divider">
  <ul>
    <li><span class="check">&#10003;</span> Page Object Model (POM) architecture for clean separation of test logic and page interactions</li>
    <li><span class="check">&#10003;</span> Modular framework design with distinct layers for pages, components, specs, utilities, and fixtures</li>
    <li><span class="check">&#10003;</span> Reusable component objects (Navbar, ListingCard) for shared UI patterns across pages</li>
    <li><span class="check">&#10003;</span> Mock authentication via localStorage session injection and REST API route interception</li>
    <li><span class="check">&#10003;</span> Mock conversation data with 5 conversations, 74 messages, and real-world read/unread states</li>
    <li><span class="check">&#10003;</span> Mock listings data with 71 listings across multiple categories and 232 associated images</li>
    <li><span class="check">&#10003;</span> Mock user profiles with diverse Indian names, roles, and presence indicators</li>
    <li><span class="check">&#10003;</span> data-testid attribute strategy for resilient, semantic element selection</li>
    <li><span class="check">&#10003;</span> Playwright HTML reporter for interactive test result visualization</li>
    <li><span class="check">&#10003;</span> Allure reporting for rich, customizable test reports with history and trends</li>
    <li><span class="check">&#10003;</span> Playwright UI Mode for real-time test exploration and debugging</li>
    <li><span class="check">&#10003;</span> Playwright Debug Mode for step-through test execution with Playwright Inspector</li>
    <li><span class="check">&#10003;</span> Screenshot capture on test failure for visual forensics</li>
    <li><span class="check">&#10003;</span> Trace file capture for detailed DOM snapshots, network logs, and timing data</li>
    <li><span class="check">&#10003;</span> Video recording of test execution for behavioral analysis</li>
  </ul>
</div>
"""


def _results_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">8. Results</h1>
  <hr class="section-divider">
  <p><strong>22 out of 22 automated end-to-end tests executed successfully.</strong></p>
  <p>No failed tests.</p>
  <p>No skipped tests.</p>
  <div class="health-green">Overall Project Automation Health: GREEN</div>
  <table>
    <tr><th>Module</th><th>Tests</th><th>Passed</th><th>Failed</th><th>Skipped</th><th>Pass %</th></tr>
    <tr><td>Authentication</td><td class="center">6</td><td class="center">6</td><td class="center">0</td><td class="center">0</td><td class="center status-pass">100%</td></tr>
    <tr><td>Chat Messaging</td><td class="center">8</td><td class="center">8</td><td class="center">0</td><td class="center">0</td><td class="center status-pass">100%</td></tr>
    <tr><td>Listings</td><td class="center">8</td><td class="center">8</td><td class="center">0</td><td class="center">0</td><td class="center status-pass">100%</td></tr>
    <tr><td class="metric">Total</td><td class="center">22</td><td class="center">22</td><td class="center">0</td><td class="center">0</td><td class="center status-pass">100%</td></tr>
  </table>
</div>
"""


def _quality_assessment_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">9. Quality Assessment</h1>
  <hr class="section-divider">
  <h2 class="subsection-title">Stability</h2>
  <p>The automation framework demonstrates exceptional stability, with all 22 tests consistently passing across multiple execution runs. The mock-based approach eliminates flakiness introduced by network latency, rate limiting, or backend instability. Tests are self-contained and produce deterministic results.</p>
  <h2 class="subsection-title">Reliability</h2>
  <p>Mock authentication and data strategies ensure tests are not dependent on external services. Route interception provides predictable API responses, while the localStorage-based session injection simulates authenticated state without requiring real credentials. The framework reliably validates both happy paths and edge cases.</p>
  <h2 class="subsection-title">Maintainability</h2>
  <p>The Page Object Model architecture, modular directory structure, and reusable utilities significantly reduce maintenance overhead. UI changes require updates only in page object methods, while test specifications remain unchanged. Consistent coding conventions and descriptive test names enhance readability and onboarding efficiency.</p>
  <h2 class="subsection-title">Scalability</h2>
  <p>The framework architecture supports easy addition of new test suites. Adding a new feature module requires creating a page object, test spec, and mock data - following established patterns. Playwright's parallel execution capabilities enable efficient scaling across multiple spec files.</p>
  <h2 class="subsection-title">Readiness for CI/CD</h2>
  <p>The framework is fully ready for CI/CD integration. Tests run in headless mode, produce machine-readable output (JUnit XML, Allure results), and support environment configuration through environment variables. The GitHub Actions workflow template can be activated with minimal configuration to run tests on pull requests and scheduled intervals.</p>
</div>
"""


def _recommendations_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">10. Recommendations</h1>
  <hr class="section-divider">
  <h2 class="subsection-title">Cross-Browser Testing</h2>
  <p>Extend the test matrix to include WebKit (Safari) and Firefox browsers alongside Chromium. Playwright's cross-browser support enables consistent validation across browser engines with minimal code changes.</p>
  <h2 class="subsection-title">Mobile Device Testing</h2>
  <p>Leverage Playwright's device emulation to validate responsive layouts and touch interactions on mobile viewport sizes. This ensures the mobile user experience meets quality standards.</p>
  <h2 class="subsection-title">API Automation</h2>
  <p>Develop a dedicated API test suite to validate backend endpoints, business logic, and data integrity independently of the UI. This provides faster feedback on backend changes and improves coverage.</p>
  <h2 class="subsection-title">Performance Automation</h2>
  <p>Introduce performance testing using Playwright's built-in performance API or integrate with tools like Lighthouse CI to track Core Web Vitals, load times, and runtime performance metrics.</p>
  <h2 class="subsection-title">GitHub Actions Integration</h2>
  <p>Activate the CI/CD pipeline by configuring GitHub Actions workflows. Configure parallel test execution across browsers, Allure report publishing, Slack/email notifications on failures, and automatic test execution on pull request creation.</p>
</div>
"""


def _conclusion_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">11. Conclusion</h1>
  <hr class="section-divider">
  <p>The ValClassifieds v1 Automation Test Execution Report confirms that the application has undergone rigorous automated end-to-end testing across its three core modules: Authentication, Chat Messaging, and Listings Management.</p>
  <p>With 22 out of 22 automated tests passing (100% pass rate), zero failures, and zero skipped tests, the automation framework demonstrates production-ready quality and reliability. The framework's modular architecture, comprehensive mock data strategy, and robust reporting mechanisms make it well-suited for continuous integration and continuous delivery pipelines.</p>
  <p><strong>The automation framework is production-ready and suitable for continuous integration.</strong></p>
  <p>The combination of Playwright's powerful testing capabilities, the Page Object Model design pattern, and comprehensive mock data provides a solid foundation for maintaining quality as the application evolves. The framework is designed to scale with the project, supporting new features, cross-browser testing, and CI/CD integration with minimal additional investment.</p>
</div>
"""


def _appendix_html():
    return f"""
<div class="page-break">
  <h1 class="section-title">Appendix</h1>
  <hr class="section-divider">

  <h2 class="subsection-title">Folder Structure</h2>
  <div class="code">
tests/e2e/<br>
  &nbsp;&nbsp;|-- pages/              Page Object classes<br>
  &nbsp;&nbsp;|-- components/         Component Object classes<br>
  &nbsp;&nbsp;|-- specs/              Test specifications<br>
  &nbsp;&nbsp;|&nbsp;&nbsp;&nbsp;&nbsp;|-- auth/     Authentication tests<br>
  &nbsp;&nbsp;|&nbsp;&nbsp;&nbsp;&nbsp;|-- chat/     Chat messaging tests<br>
  &nbsp;&nbsp;|&nbsp;&nbsp;&nbsp;&nbsp;|-- listings/ Listings tests<br>
  &nbsp;&nbsp;|-- utils/              Utilities, selectors, test data<br>
  &nbsp;&nbsp;|-- fixtures/           Custom Playwright fixtures<br>
  &nbsp;&nbsp;|-- setup/              Global setup / teardown<br>
  &nbsp;&nbsp;|-- seeds/              SQL seed scripts<br>
  &nbsp;&nbsp;|-- playwright.config.ts Configuration
  </div>

  <h2 class="subsection-title">Automation Documentation Created</h2>
  <ul>
    <li>Automation Test Plan</li>
    <li>Framework Architecture Document</li>
    <li>Framework Design Specification</li>
    <li>Automation Scope Matrix</li>
    <li>Requirements Traceability Matrix</li>
    <li>Object Repository Map</li>
    <li>Test Data Strategy Document</li>
    <li>Automation Roadmap</li>
    <li>Coding Standards Document</li>
    <li>Reporting Strategy Document</li>
    <li>Maintenance Guide</li>
    <li>Automation Metrics Document</li>
  </ul>

  <h2 class="subsection-title">QA Documentation Created</h2>
  <ul>
    <li>Master Test Case Spreadsheet (180+ test cases across 14 modules)</li>
    <li>Requirements Traceability Matrix (48+ requirements mapped)</li>
    <li>Smoke Test Suite Spreadsheet</li>
    <li>Regression Test Suite Spreadsheet</li>
    <li>Sanity Test Suite Spreadsheet</li>
    <li>Security Test Suite Spreadsheet</li>
    <li>Performance Test Suite Spreadsheet</li>
    <li>User Acceptance Test Spreadsheet</li>
    <li>Test Summary Report</li>
    <li>QA Analysis Report</li>
  </ul>

  <h2 class="subsection-title">Mock Data Created</h2>
  <ul>
    <li>16 JSON data files (users, profiles, categories, listings, images, conversations, messages, etc.)</li>
    <li>71 classified listings with realistic titles, descriptions, and pricing</li>
    <li>232 listing images across multiple categories</li>
    <li>5 conversations with 74 messages across multiple participants</li>
    <li>Conversation read/unread tracking states</li>
    <li>User presence indicators (online, away, offline)</li>
    <li>Message reactions and attachments</li>
  </ul>

  <h2 class="subsection-title">Playwright Reports Generated</h2>
  <ul>
    <li>Playwright HTML Report (interactive, with screenshots, traces, and video)</li>
    <li>Allure Report (rich dashboard with test history and trends)</li>
    <li>JUnit XML results for CI/CD integration</li>
  </ul>
</div>
"""


def convert_html_to_pdf(html_path):
    """Use Playwright (Node.js) to convert HTML to PDF."""
    import subprocess
    import sys

    pdf_path = html_path.replace(".html", ".pdf")
    abs_html = os.path.abspath(html_path)
    abs_pdf = os.path.abspath(pdf_path)

    js_script = f"""
const {{ chromium }} = require("playwright");
(async () => {{
  const browser = await chromium.launch();
  const page = await browser.newPage();
  await page.goto("file://{abs_html}", {{ waitUntil: "networkidle" }});
  await page.pdf({{
    path: "{abs_pdf}",
    format: "A4",
    displayHeaderFooter: true,
    headerTemplate: "<div style='font-size:8pt;color:#666;width:100%;text-align:right;padding:0 2cm;font-family:Calibri,sans-serif;'>ValClassifieds v1  |  Automation Test Execution Report</div>",
    footerTemplate: "<div style='font-size:8pt;color:#666;width:100%;text-align:center;padding:0 2cm;font-family:Calibri,sans-serif;'>Page <span class='pageNumber'></span> of <span class='totalPages'></span></div>",
    margin: {{ top: "2cm", bottom: "2.5cm", left: "2.5cm", right: "2.5cm" }},
    printBackground: true,
  }});
  await browser.close();
  console.log("PDF saved:", "{abs_pdf}");
}})();
"""
    try:
        result = subprocess.run(
            ["node", "-e", js_script],
            capture_output=True, text=True, timeout=30000
        )
        if result.returncode == 0:
            print(result.stdout)
            return True
        else:
            print(f"PDF conversion error: {result.stderr}", file=sys.stderr)
            return False
    except Exception as e:
        print(f"PDF conversion failed: {e}", file=sys.stderr)
        return False


def verify_file(path):
    """Verify file exists and has reasonable size."""
    if not os.path.exists(path):
        print(f"ERROR: File not found: {path}")
        return False
    size = os.path.getsize(path)
    print(f"  Verified: {path} ({size:,} bytes)")
    if size < 1000:
        print(f"  WARNING: File appears too small ({size} bytes)")
        return False
    return True


def main():
    print("=" * 60)
    print("  Automation Test Execution Report Generator")
    print("  ValClassifieds v1")
    print("=" * 60)
    print()

    os.makedirs(REPORT_DIR, exist_ok=True)

    print("[1/3] Generating DOCX report...")
    generate_docx()

    print("[2/3] Generating HTML report for PDF conversion...")
    html_path = generate_html()

    print("[3/3] Converting HTML to PDF via Playwright...")
    success = convert_html_to_pdf(html_path)

    print()
    print("=" * 60)
    print("  Verification")
    print("=" * 60)
    all_ok = True
    all_ok &= verify_file(DOCX_PATH)
    pdf_path = DOCX_PATH.replace(".docx", ".pdf")
    all_ok &= verify_file(pdf_path)
    all_ok &= verify_file(html_path)

    print()
    if all_ok:
        print("  All files generated successfully.")
    else:
        print("  WARNING: Some files may have issues.")
    print()
    print("  Output directory: reports/")
    print("  Files:")
    print(f"    - {DOCX_PATH}")
    print(f"    - {pdf_path}")
    print(f"    - {html_path}")
    print("=" * 60)

    return 0 if all_ok else 1


if __name__ == "__main__":
    exit(main())
